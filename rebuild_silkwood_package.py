from pathlib import Path
import json

import pdfplumber
from PIL import Image, ImageChops, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE = Path('Silkwood visual identity design.pdf')
ROOT = Path('Silkwood_visual_identity_extracted')
BRAND = ROOT / 'assets' / 'brand'
ROOT.mkdir(exist_ok=True)
BRAND.mkdir(parents=True, exist_ok=True)
COLORS = {
    'warm-copper': '#994E14',
    'golden-ochre': '#C79242',
    'peach-cream': '#FFE3BA',
    'vanilla-cream': '#FFF1DB',
}


def rgb(value):
    value = value.lstrip('#')
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def vector_crop(page, box, background):
    image = page.crop(box).to_image(resolution=600).original.convert('RGBA')
    matte = Image.new('RGBA', image.size, background + (255,))
    delta = ImageChops.difference(image, matte).convert('L')
    alpha = delta.point(lambda value: 0 if value < 7 else min(255, (value - 7) * 4))
    image.putalpha(alpha)
    bounds = alpha.getbbox()
    if not bounds:
        raise RuntimeError(f'No artwork detected in crop {box}')
    left, top, right, bottom = bounds
    padding = 36
    return image.crop((max(0, left-padding), max(0, top-padding), min(image.width, right+padding), min(image.height, bottom+padding)))


def png(image, name):
    image.save(BRAND / f'{name}.png', 'PNG', optimize=True)


def jpg(image, name, background):
    rendered = Image.new('RGB', image.size, background)
    rendered.paste(image, mask=image.getchannel('A'))
    rendered.save(BRAND / f'{name}.jpg', 'JPEG', quality=95, subsampling=0, optimize=True)


def recolor(image, color):
    result = Image.new('RGBA', image.size, color + (0,))
    result.putalpha(image.getchannel('A'))
    return result


with pdfplumber.open(SOURCE) as pdf:
    page = pdf.pages[0]
    light_logo = vector_crop(page, (118, 58, 318, 178), rgb(COLORS['vanilla-cream']))
    dark_logo = vector_crop(page, (70, 440, 365, 540), rgb(COLORS['warm-copper']))
    # The source's Peach Cream monogram panel, cropped inside the tile so it can be transparent.
    mark = vector_crop(page, (292, 565, 350, 625), rgb(COLORS['peach-cream']))

png(light_logo, 'logo-on-light')
jpg(light_logo, 'logo-on-light', rgb(COLORS['vanilla-cream']))
png(light_logo, 'logo-full')
jpg(light_logo, 'logo-full', rgb(COLORS['vanilla-cream']))
png(dark_logo, 'logo-on-dark')
jpg(dark_logo, 'logo-on-dark', rgb(COLORS['warm-copper']))
png(mark, 'logo-mark')
jpg(mark, 'logo-mark', rgb(COLORS['peach-cream']))

light_mark = recolor(mark, rgb(COLORS['vanilla-cream']))
png(light_mark, 'logo-mark-on-copper')
jpg(light_mark, 'logo-mark-on-copper', rgb(COLORS['warm-copper']))

for size, name in [(512, 'favicon'), (180, 'apple-touch-icon'), (32, 'favicon-32'), (16, 'favicon-16')]:
    icon = Image.new('RGBA', (size, size), rgb(COLORS['warm-copper']) + (255,))
    glyph = light_mark.copy()
    glyph.thumbnail((int(size * .58), int(size * .58)), Image.Resampling.LANCZOS)
    icon.alpha_composite(glyph, ((size-glyph.width)//2, (size-glyph.height)//2))
    icon.save(BRAND / f'{name}.png', 'PNG', optimize=True)
Image.open(BRAND / 'favicon.png').save(BRAND / 'favicon.ico', format='ICO', sizes=[(16,16),(32,32),(48,48),(64,64),(180,180)])

palette = Image.new('RGB', (1600, 420), rgb(COLORS['vanilla-cream']))
draw = ImageDraw.Draw(palette)
for i, (name, value) in enumerate(COLORS.items()):
    x = i * 400
    draw.rectangle((x, 0, x + 400, 420), fill=rgb(value))
    ink = rgb(COLORS['vanilla-cream']) if i < 2 else rgb(COLORS['warm-copper'])
    draw.text((x + 32, 300), name.replace('-', ' ').title(), fill=ink)
    draw.text((x + 32, 340), value, fill=ink)
palette.save(BRAND / 'color-palette.png', 'PNG', optimize=True)
(BRAND / 'colors.json').write_text(json.dumps(COLORS, indent=2) + '\n', encoding='utf-8')
(BRAND / 'colors.css').write_text(':root {\n' + ''.join(f'  --silkwood-{name}: {value};\n' for name, value in COLORS.items()) + '}\n', encoding='utf-8')
(BRAND / 'README.md').write_text('''# Silkwood brand assets

- `logo-on-light.*`: copper-and-ochre logo signature for light backgrounds.
- `logo-on-dark.*`: cream-and-ochre logo signature for dark copper backgrounds.
- `logo-mark.*`: transparent SW monogram.
- `logo-mark-on-copper.*`: cream SW monogram for copper backgrounds.
- `favicon.ico` and PNG favicon files: browser/app icons.
- `colors.json`, `colors.css`, and `color-palette.png`: palette values and swatches.

PNG files preserve transparency. JPEG versions have their intended brand background.
''', encoding='utf-8')


def font(run, size=None, color=None, bold=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if bold is not None: run.bold = bold


doc = Document()
section = doc.sections[0]
for side in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(section, side, Inches(1))
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for style_name, size in [('Heading 1', 16), ('Heading 2', 13)]:
    style = doc.styles[style_name]
    style.font.name = 'Calibri'; style.font.size = Pt(size); style.font.color.rgb = RGBColor(153,78,20)
    style.paragraph_format.space_before = Pt(14); style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(5)
r = title.add_run('Silkwood Visual Identity Design')
font(r, 22, (153,78,20), True)
sub = doc.add_paragraph()
r = sub.add_run('Extracted content from the supplied visual-identity PDF')
font(r, 10, (90,90,90))
doc.add_picture(str(BRAND / 'logo-on-light.png'), width=Inches(2.7))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

content = {
    'Logo Brief': ["The Silkwood Hotel logo is designed as a refined visual signature that embodies modern luxury, quiet confidence, and timeless sophistication. It balances minimalism with character, ensuring versatility across physical and digital touchpoints. Its purpose is both identification and positioning: to immediately communicate Silkwood as a premium hospitality brand within Ikeja's upscale GRA environment."],
    'About Project': ["Silkwood Hotel is envisioned as a refined expression of modern luxury for a discerning clientele that values comfort, privacy, and elevated experiences. Situated in GRA, Ikeja, Lagos, it serves people seeking more than accommodation: a space that reflects their taste, status, and lifestyle.", "The project goes beyond aesthetics. Every element of the identity and brand language is designed to communicate quiet confidence, elegance, warmth, and premium sophistication.", "The aim is to establish Silkwood as a distinctive presence within Ikeja's hospitality landscape: elegant yet comfortable, consistent yet full of character, and a new standard for luxury hospitality."],
    'Brand Interpretation': ["Silkwood is a study in quiet luxury, where elegance is felt rather than forced. Built on refinement, balance, and intention, the brand creates an experience that feels personal, exclusive, and effortlessly elevated."],
}
for heading, paragraphs in content.items():
    doc.add_heading(heading, level=1)
    for text in paragraphs: doc.add_paragraph(text)

doc.add_paragraph('The visual identity reflects:')
for item in ['Refined simplicity and quiet confidence.', 'A grounded aesthetic that celebrates richness and calm.', 'A cohesive system designed to feel intimate and timeless.']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Logo Symbol Meaning', level=1)
for label, value in [('S', 'Represents silk: smoothness, softness, and luxury.'), ('W', 'Represents wood: strength, warmth, and stability.'), ('Flowing curves', 'Represent silk in motion, conveying elegance and grace.'), ('Overall form', 'Creates a sense of shelter, trust, and exclusivity.')]:
    p = doc.add_paragraph(); a = p.add_run(label + ': '); font(a, bold=True); b = p.add_run(value); font(b)
doc.add_heading('Colour Palette', level=1)
for name, value in COLORS.items(): doc.add_paragraph(f'{name.replace("-", " ").title()} - {value}')
doc.add_heading('Typography', level=1)
doc.add_paragraph('Primary Typeface: Trajan Pro')
doc.add_paragraph('Secondary Typeface: AventaVariable')
doc.add_heading('Tagline', level=1)
doc.add_paragraph('Refined Living. Elevated Stay.')
doc.save(ROOT / 'Silkwood_visual_identity_content.docx')

print(ROOT)
